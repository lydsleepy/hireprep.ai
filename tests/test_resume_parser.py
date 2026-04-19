import pytest
from app.resume_parser import extract_resume_text


def test_raises_on_doc_format():
    with pytest.raises(ValueError, match="Legacy .doc format"):
        extract_resume_text(b"fake content", "resume.doc")


def test_raises_on_unsupported_extension():
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_resume_text(b"fake content", "resume.txt")


def test_raises_on_empty_pdf():
    # A valid but text-empty PDF (minimal PDF structure)
    minimal_pdf = b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog >>\nendobj\nxref\n0 1\n0000000000 65535 f \ntrailer\n<< /Size 1 /Root 1 0 R >>\nstartxref\n9\n%%EOF"
    with pytest.raises(ValueError, match="couldn't read any text"):
        extract_resume_text(minimal_pdf, "empty.pdf")


def test_raises_on_short_text_pdf(tmp_path):
    from unittest.mock import patch, MagicMock
    mock_reader = MagicMock()
    mock_page = MagicMock()
    mock_page.extract_text.return_value = "Hi"
    mock_reader.pages = [mock_page]
    with patch("app.resume_parser.PdfReader", return_value=mock_reader):
        with pytest.raises(ValueError, match="couldn't read any text"):
            extract_resume_text(b"fakepdf", "short.pdf")


def test_pdf_extraction_returns_stripped_text():
    from unittest.mock import patch, MagicMock
    mock_reader = MagicMock()
    mock_page = MagicMock()
    mock_page.extract_text.return_value = "  John Doe\n\n\nSoftware Engineer at Acme Corp\nExperienced developer with 5 years building Python applications and distributed systems. Led migration from monolith to microservices.\n\n"
    mock_reader.pages = [mock_page]
    with patch("app.resume_parser.PdfReader", return_value=mock_reader):
        result = extract_resume_text(b"fakepdf", "resume.pdf")
    assert "John Doe" in result
    assert "Software Engineer" in result
    # Should not have more than one consecutive blank line
    assert "\n\n\n" not in result


def test_docx_extraction_includes_table_text():
    from unittest.mock import patch, MagicMock
    mock_doc = MagicMock()
    mock_para = MagicMock()
    mock_para.text = "John Doe - Senior Software Engineer with extensive experience in Python and distributed systems"
    mock_cell = MagicMock()
    mock_cell.text = "Python | JavaScript | SQL | Docker | Kubernetes | React | PostgreSQL"
    mock_row = MagicMock()
    mock_row.cells = [mock_cell]
    mock_table = MagicMock()
    mock_table.rows = [mock_row]
    mock_doc.paragraphs = [mock_para]
    mock_doc.tables = [mock_table]
    with patch("app.resume_parser.Document", return_value=mock_doc):
        result = extract_resume_text(b"fakedocx", "resume.docx")
    assert "John Doe" in result
    assert "Python" in result


def test_invalid_docx_bytes_produces_friendly_error():
    """Corrupted DOCX bytes should produce a user-friendly ValueError."""
    with pytest.raises(ValueError, match="couldn't read any text"):
        extract_resume_text(b"not a valid docx file at all", "corrupt.docx")


def test_real_docx_short_text_rejected():
    """Real python-docx generated DOCX with short text should be rejected."""
    from docx import Document as DocxDocument
    import io

    doc = DocxDocument()
    doc.add_paragraph("Short")
    buf = io.BytesIO()
    doc.save(buf)

    with pytest.raises(ValueError, match="couldn't read any text"):
        extract_resume_text(buf.getvalue(), "short.docx")
