import io

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT

from docx import Document as DocxDocument
from docx.shared import Inches, Pt


def to_pdf(text: str, title: str) -> bytes:
    """Render plain text to a PDF and return the bytes."""
    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=title,
    )

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "body",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=17,
        spaceAfter=6,
        alignment=TA_LEFT,
    )

    story = []
    paragraphs = text.split("\n\n") if text else [""]

    for para in paragraphs:
        para = para.strip()
        if not para:
            story.append(Spacer(1, 6))
            continue

        # Escape XML special chars that reportlab chokes on
        para = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        # Preserve line breaks within a paragraph using <br/>
        para = para.replace("\n", "<br/>")
        story.append(Paragraph(para, body_style))

    doc.build(story)
    return buf.getvalue()


def to_docx(text: str, title: str) -> bytes:
    """Render plain text to a DOCX and return the bytes."""
    document = DocxDocument()

    # Set 1-inch margins on all sections
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = text.split("\n") if text else [""]

    for line in lines:
        stripped = line.strip()
        para = document.add_paragraph()
        run = para.add_run(stripped)

        # Bold lines that look like section headers:
        # ALL CAPS lines or lines that are just a label ending with ":"
        if stripped and (stripped == stripped.upper() or stripped.endswith(":")):
            run.bold = True

        run.font.name = "Calibri"
        run.font.size = Pt(11)

        # Minimal spacing
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
