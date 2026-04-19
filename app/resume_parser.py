import io
import re
from pathlib import Path

from pypdf import PdfReader, errors as pypdf_errors
from docx import Document


_MIN_TEXT_LENGTH = 20


def extract_resume_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from a PDF or DOCX resume.

    Raises ValueError with user-facing messages on unsupported formats,
    unreadable files, or files with too little extracted text.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".doc":
        raise ValueError(
            "Legacy .doc format isn't supported. Please save your resume as "
            ".docx or .pdf and try again."
        )
    if suffix == ".pdf":
        return _extract_pdf(file_bytes)
    if suffix == ".docx":
        return _extract_docx(file_bytes)

    raise ValueError(
        f"Unsupported file type '{suffix}'. Please upload a .pdf or .docx file."
    )


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
    except pypdf_errors.PdfReadError:
        raise ValueError(
            "We couldn't read any text from this file. If it's a scanned image, "
            "please upload a text-based PDF or .docx instead."
        )
    raw = "\n".join(pages)
    text = _clean(raw)
    _require_min_text(text)
    return text


def _extract_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))

    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Resumes often use tables for layout — extract those too
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append("  ".join(row_texts))

    text = _clean("\n".join(parts))
    _require_min_text(text)
    return text


def _clean(text: str) -> str:
    """Collapse excessive blank lines and strip leading/trailing whitespace."""
    text = text.strip()
    # Replace 3+ consecutive newlines with 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def _require_min_text(text: str) -> None:
    if len(text) < _MIN_TEXT_LENGTH:
        raise ValueError(
            "We couldn't read any text from this file. If it's a scanned image, "
            "please upload a text-based PDF or .docx instead."
        )
